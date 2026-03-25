"""
Module de conversion de fichiers pour Albert IA Agentic.
Gère les images, PDF, DOCX, XLS et autres formats.
"""

import base64
import io
import os
import uuid
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass, field
from enum import Enum
import re

# Limites de configuration
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
MAX_FILES = 10
MAX_IMAGE_DIMENSION = 2048  # pixels


class FileType(Enum):
    IMAGE = "image"
    PDF = "pdf"
    DOCX = "docx"
    XLSX = "xlsx"
    TEXT = "text"
    UNKNOWN = "unknown"


@dataclass
class Attachment:
    """Représente un fichier attaché."""

    id: str
    name: str
    file_type: FileType
    size: int
    mime_type: str
    content: str = ""  # base64 encoded
    text_content: str = ""  # Pour DOCX/XLS
    pages: List[str] = field(default_factory=list)  # Pour PDF


def get_file_type(filename: str, mime_type: str = "") -> FileType:
    """Détermine le type de fichier selon extension et mime type."""
    ext = Path(filename).suffix.lower()
    mime = mime_type.lower()

    # Images
    image_exts = {".jpg", ".jpeg", ".png", ".gif", ".webp", ".bmp"}
    image_mimes = {"image/jpeg", "image/png", "image/gif", "image/webp", "image/bmp"}

    if ext in image_exts or mime in image_mimes:
        return FileType.IMAGE

    # PDF
    if ext == ".pdf" or mime == "application/pdf":
        return FileType.PDF

    # Word
    if ext in {".docx", ".doc"} or "word" in mime:
        return FileType.DOCX

    # Excel
    if ext in {".xlsx", ".xls"} or "excel" in mime or "spreadsheet" in mime:
        return FileType.XLSX

    # Text
    text_exts = {".txt", ".md", ".csv", ".json", ".xml", ".html", ".py", ".js", ".ts"}
    text_mimes = {"text/", "application/json", "application/xml"}

    if ext in text_exts or any(mime.startswith(m) for m in text_mimes):
        return FileType.TEXT

    return FileType.UNKNOWN


def resize_image(image_data: bytes, max_dimension: int = MAX_IMAGE_DIMENSION) -> bytes:
    """Redimensionne une image si nécessaire."""
    try:
        from PIL import Image
        import PIL.Image

        img = Image.open(io.BytesIO(image_data))

        # Convertir en RGB si nécessaire
        if img.mode in ("RGBA", "P"):
            img = img.convert("RGB")

        # Calculer les nouvelles dimensions
        width, height = img.size
        if width > max_dimension or height > max_dimension:
            ratio = min(max_dimension / width, max_dimension / height)
            new_width = int(width * ratio)
            new_height = int(height * ratio)
            img = img.resize((new_width, new_height), PIL.Image.Resampling.LANCZOS)

        # Réencoder en JPEG optimisé
        output = io.BytesIO()
        img.save(output, format="JPEG", quality=85, optimize=True)
        return output.getvalue()

    except ImportError:
        # Si PIL non disponible, retourner l'image originale
        return image_data


def extract_text_from_docx(file_content: bytes) -> Tuple[str, List[str]]:
    """Extrait le texte et les images d'un fichier DOCX."""
    try:
        from docx import Document
        from docx.oxml.ns import qn
        import zipfile

        doc = Document(io.BytesIO(file_content))

        # Extraire le texte
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)

        # Extraire le texte des tableaux
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                if any(row_text):
                    full_text.append(" | ".join(row_text))

        # Extraire les images
        images = []
        try:
            with zipfile.ZipFile(io.BytesIO(file_content)) as zf:
                for name in zf.namelist():
                    if name.startswith("word/media/") and name.endswith(
                        (".png", ".jpg", ".jpeg", ".gif", ".webp")
                    ):
                        img_data = zf.read(name)
                        # Resize et convertir en base64
                        img_data = resize_image(img_data)
                        b64 = base64.b64encode(img_data).decode("utf-8")
                        images.append(b64)
        except Exception:
            pass

        return "\n\n".join(full_text), images

    except ImportError:
        return "[Impossible d'extraire le texte - python-docx non installé]", []
    except Exception as e:
        return f"[Erreur d'extraction DOCX: {str(e)}]", []


def extract_text_from_xlsx(file_content: bytes) -> Tuple[str, List[str]]:
    """Extrait le texte d'un fichier Excel."""
    try:
        import openpyxl
        from openpyxl.drawing.image import Image as XLImage

        wb = openpyxl.load_workbook(io.BytesIO(file_content), data_only=True)

        full_text = []
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            full_text.append(f"\n## Feuille: {sheet_name}\n")

            for row in sheet.iter_rows(values_only=True):
                row_text = []
                for cell in row:
                    if cell is not None:
                        cell_str = str(cell)
                        if cell_str.strip():
                            row_text.append(cell_str.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))

        return "\n".join(full_text), []

    except ImportError:
        return "[Impossible d'extraire le texte - openpyxl non installé]", []
    except Exception as e:
        return f"[Erreur d'extraction XLSX: {str(e)}]", []


def convert_pdf_to_images_and_text(file_content: bytes) -> Tuple[List[str], str]:
    """Convertit un PDF en images (une par page) et extrait le texte."""
    try:
        import fitz  # PyMuPDF

        pdf = fitz.open(stream=file_content, filetype="pdf")
        images = []
        full_text = []

        for page_num in range(len(pdf)):
            page = pdf[page_num]
            
            # Extraction du texte
            text = page.get_text().strip()
            if text:
                full_text.append(f"--- PAGE {page_num + 1} ---\n{text}")

            # Rendu en image (200 DPI pour une meilleure qualité OCR)
            mat = fitz.Matrix(200 / 72, 200 / 72)
            pix = page.get_pixmap(matrix=mat)
            img_data = pix.tobytes("jpeg")

            # Rezise si dimension vraiment excessive (> 4096)
            img_data = resize_image(img_data, max_dimension=4096)

            # Encoder en base64
            b64 = base64.b64encode(img_data).decode("utf-8")
            images.append(b64)

        pdf.close()
        return images, "\n\n".join(full_text)

    except ImportError:
        return [], ""
    except Exception:
        return [], ""


def process_file(
    file_content: bytes, filename: str, mime_type: str = ""
) -> Optional[Attachment]:
    """
    Traite un fichier et retourne un Attachment avec le contenu extrait.
    """
    file_size = len(file_content)

    # Vérifier la taille
    if file_size > MAX_FILE_SIZE:
        return None

    file_type = get_file_type(filename, mime_type)
    attachment_id = str(uuid.uuid4())

    attachment = Attachment(
        id=attachment_id,
        name=filename,
        file_type=file_type,
        size=file_size,
        mime_type=mime_type or "application/octet-stream",
    )

    if file_type == FileType.IMAGE:
        # Resize et encoder
        processed = resize_image(file_content)
        attachment.content = base64.b64encode(processed).decode("utf-8")

    elif file_type == FileType.PDF:
        # Convertir chaque page en image et extraire le texte
        images, text = convert_pdf_to_images_and_text(file_content)
        attachment.pages = images
        attachment.text_content = text
        if images:
            attachment.content = images[0]  # Première page en preview

    elif file_type == FileType.DOCX:
        text, images = extract_text_from_docx(file_content)
        attachment.text_content = text
        attachment.pages = images
        if images:
            attachment.content = images[0]

    elif file_type == FileType.XLSX:
        text, images = extract_text_from_xlsx(file_content)
        attachment.text_content = text

    elif file_type == FileType.TEXT:
        try:
            attachment.text_content = file_content.decode("utf-8")
        except:
            try:
                attachment.text_content = file_content.decode("latin-1")
            except:
                attachment.text_content = "[Impossible de lire le fichier]"
    else:
        # Fichier non supporté pour le moment
        return None

    return attachment


def get_mime_type_from_ext(filename: str) -> str:
    """Retourne le MIME type basé sur l'extension."""
    ext = Path(filename).suffix.lower()
    mime_map = {
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".png": "image/png",
        ".gif": "image/gif",
        ".webp": "image/webp",
        ".bmp": "image/bmp",
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".doc": "application/msword",
        ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ".xls": "application/vnd.ms-excel",
        ".txt": "text/plain",
        ".md": "text/markdown",
        ".csv": "text/csv",
        ".json": "application/json",
    }
    return mime_map.get(ext, "application/octet-stream")


def attachment_to_content_blocks(attachment: Attachment) -> List[Dict[str, Any]]:
    """
    Convertit un Attachment en liste de content blocks pour l'API Mistral.
    """
    blocks = []

    # Texte extrait d'abord (pour PDF, DOCX, XLSX, etc.)
    if attachment.text_content and attachment.file_type in {
        FileType.PDF,
        FileType.DOCX,
        FileType.XLSX,
        FileType.TEXT,
    }:
        content = f"Fichier: {attachment.name}\n"
        if attachment.file_type == FileType.PDF:
             content += "Contenu texte extrait du PDF :\n"
        
        content += f"\n{attachment.text_content[:50000]}"
        
        blocks.append(
            {
                "type": "text",
                "text": content,
            }
        )

    # Images (pages pour PDF, images pour DOCX)
    if attachment.pages:
        for page_b64 in attachment.pages[:10]:  # Max 10 pages
            blocks.append(
                {
                    "type": "image_url",
                    "image_url": {"url": f"data:image/jpeg;base64,{page_b64}"},
                }
            )
    elif attachment.content:
        # Image unique
        mime = "image/jpeg"
        if attachment.mime_type.startswith("image/"):
            mime = attachment.mime_type
        blocks.append(
            {
                "type": "image_url",
                "image_url": {"url": f"data:{mime};base64,{attachment.content}"},
            }
        )

    return blocks
